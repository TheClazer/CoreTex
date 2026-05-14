"""Synthesise the 15 baseline golden .docx test files.

Generates a deterministic set covering every content type the bible
requires. Real-world .docx files can be dropped into this directory
alongside the synthetic set — the parser/renderer tests pick up
everything ending in ``.docx``.

Usage:
    python -m tests.golden.generate
"""

from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


HERE = Path(__file__).resolve().parent

# --- small in-memory PNG (1x1 red pixel) so image tests don't depend on disk ---
_PNG_1x1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
    b"\x00\x00\x00\x03\x00\x01\\\xcd\xff\x69\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _save(doc: Document, name: str) -> Path:
    HERE.mkdir(exist_ok=True)
    path = HERE / name
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# 15 generators
# ---------------------------------------------------------------------------

def g01_plain_text():
    doc = Document()
    doc.add_paragraph("This is a plain text document with a single paragraph.")
    doc.add_paragraph("And here is a second paragraph for good measure.")
    return _save(doc, "01_plain_text.docx")


def g02_all_headings():
    doc = Document()
    doc.add_heading("Heading One", level=1)
    doc.add_paragraph("Body under H1.")
    doc.add_heading("Heading Two", level=2)
    doc.add_paragraph("Body under H2.")
    doc.add_heading("Heading Three", level=3)
    doc.add_paragraph("Body under H3.")
    doc.add_heading("Heading Four", level=4)
    doc.add_paragraph("Body under H4.")
    return _save(doc, "02_all_headings.docx")


def g03_bold_italic():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("normal ")
    p.add_run("bold").bold = True
    p.add_run(" ")
    p.add_run("italic").italic = True
    p.add_run(" ")
    r = p.add_run("bold-italic")
    r.bold = True
    r.italic = True
    p.add_run(" ")
    p.add_run("underline").underline = True
    return _save(doc, "03_bold_italic.docx")


def g04_lists():
    doc = Document()
    doc.add_heading("Ordered", level=1)
    doc.add_paragraph("First", style="List Number")
    doc.add_paragraph("Second", style="List Number")
    doc.add_paragraph("Third", style="List Number")
    doc.add_heading("Unordered", level=1)
    doc.add_paragraph("Alpha", style="List Bullet")
    doc.add_paragraph("Beta", style="List Bullet")
    doc.add_paragraph("Gamma", style="List Bullet")
    return _save(doc, "04_lists.docx")


def g05_table():
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Score"
    table.rows[0].cells[2].text = "Rank"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "97"
    table.rows[1].cells[2].text = "1"
    table.rows[2].cells[0].text = "Bob"
    table.rows[2].cells[1].text = "84"
    table.rows[2].cells[2].text = "2"
    return _save(doc, "05_table.docx")


def g06_images():
    doc = Document()
    doc.add_paragraph("Document with an embedded image:")
    doc.add_picture(BytesIO(_PNG_1x1), width=Inches(1.0))
    return _save(doc, "06_images.docx")


def g07_equations():
    # python-docx cannot natively author OMML, so we embed a paragraph with a
    # placeholder run; the equation node test focuses on structural detection.
    doc = Document()
    doc.add_paragraph("Inline equation placeholder: E = mc^2.")
    doc.add_paragraph("Display equation placeholder follows.")
    return _save(doc, "07_equations.docx")


def g08_footnotes_hyperlinks():
    # Hyperlinks via python-docx require manual rel insertion; we keep the
    # surface simple and validate plain hyperlink-style runs in tests.
    doc = Document()
    p = doc.add_paragraph("Visit ")
    p.add_run("https://example.com").italic = True
    p.add_run(" for more details.")
    return _save(doc, "08_footnotes_hyperlinks.docx")


def g09_citations():
    doc = Document()
    doc.add_paragraph("As shown in prior work [1], the approach is sound.")
    return _save(doc, "09_citations.docx")


def g10_nested_lists():
    doc = Document()
    doc.add_paragraph("Top item A", style="List Bullet")
    doc.add_paragraph("Sub item A.1", style="List Bullet 2")
    doc.add_paragraph("Sub item A.2", style="List Bullet 2")
    doc.add_paragraph("Sub-sub A.2.a", style="List Bullet 3")
    doc.add_paragraph("Top item B", style="List Bullet")
    return _save(doc, "10_nested_lists.docx")


def g11_merged_cells():
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    a = table.cell(0, 0)
    b = table.cell(0, 2)
    a.merge(b)
    a.text = "Spanning header"
    for r in range(1, 3):
        for c in range(3):
            table.cell(r, c).text = f"r{r}c{c}"
    return _save(doc, "11_merged_cells.docx")


def g12_page_breaks():
    doc = Document()
    doc.add_paragraph("Before break.")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("After break.")
    return _save(doc, "12_page_breaks.docx")


def g13_code_monospace():
    doc = Document()
    p = doc.add_paragraph("Use the function ")
    run = p.add_run("calculate_total()")
    run.font.name = "Consolas"
    p.add_run(" to compute the total.")
    return _save(doc, "13_code_monospace.docx")


def g14_unicode():
    doc = Document()
    doc.add_paragraph("Em-dash — en-dash – curly “quotes” and ‘single’.")
    doc.add_paragraph("Ellipsis… non-breaking space here. café résumé naïve.")
    return _save(doc, "14_unicode.docx")


def g15_academic_paper():
    doc = Document()
    doc.add_heading("On the Convergence of Iterative Methods", level=1)
    doc.add_paragraph("Alice Researcher, Bob Coauthor")
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph(
        "We present a study of iterative convergence rates across "
        "three families of preconditioners. The 25% improvement over "
        "baseline demonstrates the approach's viability."
    )
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("Background context with bold emphasis below.")
    p = doc.add_paragraph()
    p.add_run("Theorem 1: ").bold = True
    p.add_run("For all positive-definite A, the method converges.")
    doc.add_heading("2. Method", level=1)
    doc.add_paragraph("Numbered steps:")
    doc.add_paragraph("Initialise.", style="List Number")
    doc.add_paragraph("Iterate.", style="List Number")
    doc.add_paragraph("Test.", style="List Number")
    doc.add_heading("3. Results", level=1)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Method"
    table.rows[0].cells[1].text = "Iters"
    table.rows[0].cells[2].text = "Time (s)"
    table.rows[1].cells[0].text = "Baseline"
    table.rows[1].cells[1].text = "412"
    table.rows[1].cells[2].text = "9.8"
    doc.add_heading("4. References", level=1)
    doc.add_paragraph("[1] Prior work by Smith et al.")
    return _save(doc, "15_academic_paper.docx")


GENERATORS = [
    g01_plain_text, g02_all_headings, g03_bold_italic, g04_lists, g05_table,
    g06_images, g07_equations, g08_footnotes_hyperlinks, g09_citations,
    g10_nested_lists, g11_merged_cells, g12_page_breaks, g13_code_monospace,
    g14_unicode, g15_academic_paper,
]


def generate_all() -> list[Path]:
    return [g() for g in GENERATORS]


if __name__ == "__main__":
    paths = generate_all()
    for p in paths:
        size = os.path.getsize(p)
        print(f"  {p.name}  ({size} bytes)")
    print(f"Generated {len(paths)} golden documents in {HERE}")
