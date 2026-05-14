"""Generate a realistic academic-paper .docx for end-to-end testing.

Run:
    python -m tests.golden.make_demo
Output:
    tests/golden/DEMO_paper.docx
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from PIL import Image, ImageDraw

OUT = Path(__file__).resolve().parent / "DEMO_paper.docx"


def _figure_png() -> bytes:
    """A small heatmap-style figure so the conversion actually has an image."""
    img = Image.new("RGB", (320, 200), (250, 250, 252))
    draw = ImageDraw.Draw(img)
    palette = [
        (37, 99, 235), (59, 130, 246), (96, 165, 250),
        (147, 197, 253), (252, 165, 165), (239, 68, 68),
        (220, 38, 38), (185, 28, 28),
    ]
    for r in range(8):
        for c in range(8):
            color = palette[(r + c) % len(palette)]
            draw.rectangle([20 + c * 35, 20 + r * 20, 20 + (c + 1) * 35, 20 + (r + 1) * 20], fill=color)
    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


_PNG = _figure_png()


def build():
    doc = Document()

    # ── Title (H1) ──────────────────────────────────────────────────
    doc.add_heading("Neural Compression of Sparse Tensors for Real-Time Inference", level=1)

    # ── Author block ────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("Rayyan Shaikh¹, Alice Chen², Bob Müller³")
    r.italic = True
    doc.add_paragraph(
        "¹CoreTex Labs, Bangalore  ·  ²Stanford University  ·  ³ETH Zürich"
    )

    # ── Abstract ────────────────────────────────────────────────────
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph(
        "We present a novel approach to compressing sparse activation tensors "
        "in transformer architectures, achieving a 4.2× reduction in memory "
        "footprint with only 0.3% loss in downstream task accuracy. Our method "
        "combines block-sparse pruning with a learned quantization codebook, "
        "and we evaluate it on three benchmark suites — GLUE, SuperGLUE, and "
        "MMLU — across model sizes ranging from 350M to 70B parameters."
    )

    # ── 1. Introduction ─────────────────────────────────────────────
    doc.add_heading("1. Introduction", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Modern large language models suffer from memory pressure during inference, "
        "particularly when serving multiple concurrent requests. Activation "
        "tensors — the intermediate outputs of each attention block — dominate "
        "the working set in autoregressive decoding "
    )
    p.add_run("[1]").italic = True  # placeholder citation
    p.add_run(". Prior work has explored ")
    p.add_run("static pruning").bold = True
    p.add_run(", ")
    p.add_run("dynamic sparsity").bold = True
    p.add_run(", and ")
    p.add_run("post-training quantization").bold = True
    p.add_run(", but each suffers from accuracy degradation at extreme compression ratios.")

    doc.add_paragraph(
        "Our contributions are threefold. First, we identify the "
        "“double-quantization gap” phenomenon — where naïve combination of "
        "pruning and quantization compounds errors super-linearly. Second, "
        "we propose a joint optimization objective that explicitly accounts "
        "for this interaction. Third, we release reference implementations at "
    )

    # Hyperlink-ish (python-docx doesn't have a clean hyperlink API; we use
    # plain text that the parser later treats as a run — for a real link the
    # user would mark it in Word with Ctrl+K)
    p = doc.add_paragraph()
    r = p.add_run("https://github.com/TheClazer/CoreTex")
    r.font.name = "Consolas"  # monospace — will become \texttt

    # ── 2. Related Work ─────────────────────────────────────────────
    doc.add_heading("2. Related Work", level=1)
    doc.add_paragraph(
        "Compression of neural networks has a long history. We focus on three lines:"
    )

    # Nested ordered list
    doc.add_paragraph("Structured pruning approaches.", style="List Number")
    doc.add_paragraph("Magnitude-based.", style="List Number 2")
    doc.add_paragraph("Importance-score based (Taylor expansion).", style="List Number 2")
    doc.add_paragraph("Quantization approaches.", style="List Number")
    doc.add_paragraph("Linear: GPTQ, AWQ.", style="List Number 2")
    doc.add_paragraph("Non-linear: K-means codebook, vector quantization.", style="List Number 2")
    doc.add_paragraph("Hybrid approaches (closest to ours).", style="List Number")

    # ── 3. Method ───────────────────────────────────────────────────
    doc.add_heading("3. Method", level=1)
    doc.add_heading("3.1 Notation", level=2)

    p = doc.add_paragraph()
    p.add_run("Let X ∈ ℝ").italic = True
    p.add_run("ᵇˣⁿˣᵈ").italic = True
    p.add_run(
        " denote an activation tensor with batch dimension b, sequence length n, "
        "and hidden size d. We seek a compressed representation X̃ such that "
        "‖X − X̃‖² is bounded, while the encoded form requires fewer than "
    )
    p.add_run("kd").bold = True
    p.add_run(" bits per token, for k < 8.")

    doc.add_heading("3.2 Algorithm", level=2)
    doc.add_paragraph("Our compression pipeline operates in three stages:", style="List Bullet")
    doc.add_paragraph("Block-sparse mask selection via top-k magnitude.", style="List Bullet")
    doc.add_paragraph("Codebook learning over surviving entries (Lloyd–Max).", style="List Bullet")
    doc.add_paragraph("Joint fine-tuning with the straight-through estimator.", style="List Bullet")

    # ── 4. Results (with table) ─────────────────────────────────────
    doc.add_heading("4. Experimental Results", level=1)
    doc.add_paragraph(
        "Table 1 summarises the trade-off between compression ratio and "
        "accuracy on the GLUE benchmark using the Llama-3 8B base model."
    )

    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    hdr.cells[0].text = "Method"
    hdr.cells[1].text = "Bits/param"
    hdr.cells[2].text = "GLUE avg ↑"
    hdr.cells[3].text = "Speedup ×"

    rows = [
        ("Baseline (FP16)", "16.0", "84.2", "1.00"),
        ("GPTQ-4bit", "4.0", "82.7", "2.85"),
        ("AWQ-4bit", "4.0", "83.1", "2.91"),
        ("Ours (k=3.8)", "3.8", "83.9", "3.42"),
    ]
    for i, (m, b, g, s) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = m
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = g
        table.rows[i].cells[3].text = s

    # Caption
    cap = doc.add_paragraph()
    r = cap.add_run("Table 1.")
    r.bold = True
    cap.add_run(
        " Compression vs accuracy on GLUE-dev. Our method achieves a better "
        "Pareto frontier than both GPTQ and AWQ at comparable bit-widths."
    )

    # ── Figure ──────────────────────────────────────────────────────
    doc.add_paragraph("Figure 1 visualises the codebook utilisation distribution.")
    doc.add_picture(BytesIO(_PNG), width=Inches(2.0))
    cap = doc.add_paragraph()
    r = cap.add_run("Figure 1.")
    r.bold = True
    cap.add_run(" Codebook entry utilisation (red: hot, blue: cold) across layers.")

    # ── 5. Discussion ───────────────────────────────────────────────
    doc.add_heading("5. Discussion", level=1)
    doc.add_paragraph(
        "We observed three failure modes during development. First, "
        "extreme low-bit regimes (k < 2) collapse activations onto only a "
        "handful of codebook entries — the so-called “codebook collapse” "
        "issue, well known in vector quantization literature. Second, "
        "fine-tuning instability emerges when the straight-through estimator "
        "is combined with mixed-precision training; we resolved this by "
        "clipping gradients at the codebook boundary. Third, on very long "
        "sequences (≥ 32k tokens), the block-sparse mask becomes too "
        "fine-grained and the overhead of mask storage dominates."
    )

    # Page break before References
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ── References ──────────────────────────────────────────────────
    doc.add_heading("References", level=1)
    refs = [
        "[1] Vaswani et al., “Attention Is All You Need,” NeurIPS 2017.",
        "[2] Frantar et al., “GPTQ: Accurate Post-Training Quantization,” ICLR 2023.",
        "[3] Lin et al., “AWQ: Activation-aware Weight Quantization,” MLSys 2024.",
        "[4] Han et al., “Deep Compression,” ICLR 2016.",
    ]
    for r in refs:
        doc.add_paragraph(r)

    doc.save(OUT)
    print(f"Wrote {OUT}  ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
