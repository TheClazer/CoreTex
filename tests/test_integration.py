import pytest
from httpx import AsyncClient, ASGITransport
from io import BytesIO
from unittest.mock import patch

from app.main import app
from app.config import settings
from docx import Document
from rq.exceptions import NoSuchJobError

@pytest.fixture
def minimal_docx():
    """Creates a minimal valid .docx file using python-docx and returns bytes."""
    doc = Document()
    doc.add_heading('Test Document', level=1)
    doc.add_paragraph('This is a test paragraph.')
    io_stream = BytesIO()
    doc.save(io_stream)
    return io_stream.getvalue()

@pytest.fixture
def mock_queue():
    """Mock the RQ queue to prevent actual background enqueuing during test."""
    with patch("app.api.routes.conversion_queue") as mock_q:
        yield mock_q

@pytest.fixture
def mock_routes_redis():
    """Mock the Redis connection inside routes to prevent requiring real Redis."""
    with patch("app.api.routes.redis_conn") as mock_conn:
        yield mock_conn


@pytest.mark.asyncio
async def test_root_endpoint():
    """1. Test root endpoint returns 200 with status: ok"""
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.get("/")
    assert response.status_code == 200
    assert response.json() == {"status": "ok", "service": "word-to-latex"}


@pytest.mark.asyncio
async def test_convert_rejects_non_docx():
    """2. Test POST /convert with a .txt file returns 400."""
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        files = {"file": ("test.txt", b"this is text", "text/plain")}
        response = await ac.post("/convert", files=files)
    assert response.status_code == 400
    assert "Invalid file type" in response.json()["detail"]


@pytest.mark.asyncio
async def test_convert_rejects_large_file(monkeypatch):
    """3. Test POST /convert with a file larger than MAX_FILE_SIZE_MB returns 413."""
    # Ensure settings is standard for the test
    monkeypatch.setattr(settings, "MAX_FILE_SIZE_MB", 20)
    fake_large_bytes = b"0" * (21 * 1024 * 1024)
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        files = {"file": ("large.docx", fake_large_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = await ac.post("/convert", files=files)
    assert response.status_code == 413
    assert "File too large" in response.json()["detail"]


@pytest.mark.asyncio
async def test_convert_accepts_valid_docx(minimal_docx, mock_queue):
    """4. Test POST /convert with a valid .docx enqueues the job and returns 200."""
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        files = {"file": ("test.docx", minimal_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = await ac.post("/convert", files=files)
    assert response.status_code == 200
    json_resp = response.json()
    assert "job_id" in json_resp
    assert json_resp["status"] == "queued"
    mock_queue.enqueue.assert_called_once()


@pytest.mark.asyncio
async def test_status_not_found():
    """5. Test GET /status/nonexistent-id returns status: not_found."""
    with patch("app.api.routes.Job.fetch", side_effect=NoSuchJobError):
        async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
            response = await ac.get("/status/nonexistent-id")
    assert response.status_code == 200
    assert response.json() == {"job_id": "nonexistent-id", "status": "not_found"}


@pytest.mark.asyncio
async def test_temp_not_found(mock_routes_redis):
    """6. Test GET /temp/nonexistent-id returns 404."""
    # Simulate Redis returning None for the key
    mock_routes_redis.get.return_value = None
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.get("/temp/nonexistent-id")
    assert response.status_code == 404
    assert "Temp file expired" in response.json()["detail"]
