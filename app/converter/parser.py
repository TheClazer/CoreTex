"""Surgical Word OOXML parser.

Reads a .docx file (as bytes) and walks its OOXML tree, producing a typed
``IRDocument``. The parser is the only place in the system that reads Word
XML; downstream layers (renderer, handlers) consume only the IR.

Owner: P2. Reference: word_latex_bible.pdf §2 layer 2 and team distribution.
"""

from __future__ import annotations

import logging
import zipfile
from io import BytesIO
from typing import List, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree

from app.converter.ir_schema import (
    CitationNode,
    EquationNode,
    FootnoteNode,
    HeadingNode,
    HyperlinkNode,
    ImageNode,
    IRDocument,
    IRNode,
    ListItemNode,
    ListNode,
    PageBreakNode,
    ParagraphNode,
    RunNode,
    TableCellNode,
    TableNode,
    TableRowNode,
    TextAlign,
)
from app.converter.bibliography import extract_bibliography
from app.converter.run_merger import merge_runs

logger = logging.getLogger(__name__)

# Namespace map used for raw lxml queries.
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_ALIGN_MAP = {
    "left": TextAlign.LEFT,
    "start": TextAlign.LEFT,
    "center": TextAlign.CENTER,
    "centre": TextAlign.CENTER,
    "right": TextAlign.RIGHT,
    "end": TextAlign.RIGHT,
    "both": TextAlign.JUSTIFY,
    "justify": TextAlign.JUSTIFY,
}


def _alignment_from_jc(jc_elem) -> TextAlign:
    if jc_elem is None:
        return TextAlign.LEFT
    val = jc_elem.get(qn("w:val")) or ""
    return _ALIGN_MAP.get(val.lower(), TextAlign.LEFT)


def _paragraph_alignment(p: Paragraph) -> TextAlign:
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return TextAlign.LEFT
    return _alignment_from_jc(pPr.find(qn("w:jc")))


def _heading_level(p: Paragraph) -> Optional[int]:
    """Return 1..4 if ``p`` is a Heading 1..4, else None."""
    style = (p.style.name or "").strip() if p.style else ""
    if not style.lower().startswith("heading"):
        # Word also marks titles with 'Title' style; treat as H1.
        if style.lower() == "title":
            return 1
        return None
    # 'Heading 1' -> 1
    parts = style.split()
    try:
        lvl = int(parts[-1])
    except (ValueError, IndexError):
        return None
    return max(1, min(lvl, 4))


def _is_code_font(run) -> bool:
    """Detect a monospace/code run by inspecting its rFonts."""
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return False
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        return False
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        font = rFonts.get(qn(attr)) or ""
        if any(token in font.lower() for token in ("courier", "consolas", "mono", "menlo")):
            return True
    return False


def _run_text(run) -> str:
    """Collect text from a w:r element including tabs and line breaks."""
    out: list[str] = []
    for child in run._element.iter():
        tag = etree.QName(child).localname
        if tag == "t":
            out.append(child.text or "")
        elif tag == "tab":
            out.append("\t")
        elif tag == "br":
            out.append("\n")
    return "".join(out)


def _build_runs(paragraph: Paragraph) -> List[RunNode]:
    runs: List[RunNode] = []
    for r in paragraph.runs:
        text = _run_text(r)
        if not text:
            continue
        runs.append(
            RunNode(
                text=text,
                bold=bool(r.bold),
                italic=bool(r.italic),
                underline=bool(r.underline),
                code=_is_code_font(r),
            )
        )
    # v2 run-merger: collapse Word's per-word/rsid run fragmentation so the
    # renderer emits one wrapper per formatted span instead of one per run.
    return merge_runs(runs)


def _has_page_break(paragraph: Paragraph) -> bool:
    # Inline page break: <w:br w:type="page"/>
    if paragraph._p.find(f".//{qn('w:br')}[@{qn('w:type')}='page']") is not None:
        return True
    # Paragraph property: <w:pageBreakBefore/> — break BEFORE this paragraph.
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
        return True
    return False


def _paragraph_breaks_before(paragraph: Paragraph) -> bool:
    """True when this paragraph should be preceded by \\newpage."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:pageBreakBefore")) is not None


def _section_break_after(sectPr_elem) -> bool:
    """A sectPr with <w:type w:val="nextPage"/> (or missing) starts a new page."""
    if sectPr_elem is None:
        return False
    sect_type = sectPr_elem.find(qn("w:type"))
    if sect_type is None:
        # OOXML default for section break is nextPage.
        return True
    val = sect_type.get(qn("w:val"))
    return val in (None, "nextPage", "oddPage", "evenPage")


# ---------------------------------------------------------------------------
# List detection
# ---------------------------------------------------------------------------

_LIST_STYLE_TOKENS = {
    # English
    "list number": (True, "number"),
    "list bullet": (False, "bullet"),
    "list paragraph": (False, "bullet"),
    # French
    "liste à puces": (False, "bullet"),
    "liste à numéros": (True, "number"),
    "liste numérotée": (True, "number"),
    # German
    "listenabsatz": (False, "bullet"),
    "aufzählung": (False, "bullet"),
    "nummerierung": (True, "number"),
    # Spanish
    "lista con viñetas": (False, "bullet"),
    "lista numerada": (True, "number"),
    "párrafo de lista": (False, "bullet"),
    # Italian
    "elenco puntato": (False, "bullet"),
    "elenco numerato": (True, "number"),
    # Portuguese
    "lista com marcadores": (False, "bullet"),
    "lista com numeração": (True, "number"),
}


def _list_info(paragraph: Paragraph) -> Optional[tuple[int, int]]:
    """Return (numId, ilvl) if paragraph is a list item, else None.

    Detection priority:
    1. ``w:numPr`` in pPr — the OOXML-native, locale-independent signal.
       Every list paragraph in every language sets this when Word saves.
    2. Style-name pattern: English defaults (``List Number``, ``List Bullet``)
       plus the most common European-language localised style names.
    3. Generic numbering-style probe — if the paragraph's style chains up
       to a base style that defines ``w:numPr`` in the styles part, treat
       it as a list. Catches custom-named styles and styles localised to
       languages not in the table above.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            numId_el = numPr.find(qn("w:numId"))
            ilvl_el = numPr.find(qn("w:ilvl"))
            if numId_el is not None:
                num_id = int(numId_el.get(qn("w:val"), "0"))
                ilvl = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
                return num_id, ilvl

    # Style-name table — known localisations.
    style = (paragraph.style.name or "") if paragraph.style else ""
    sl = style.lower()
    for token, (ordered, _kind) in _LIST_STYLE_TOKENS.items():
        if sl.startswith(token):
            base = 1000 if ordered else 2000
            # Indent level from trailing digit, if present.
            tail = sl[len(token):].strip()
            try:
                ilvl = int(tail) - 1 if tail else 0
            except ValueError:
                ilvl = 0
            return base, max(0, ilvl)

    # Last-resort probe: walk the style chain via the styles part to find an
    # ancestor that defines numPr. Locale-independent.
    if paragraph.style is not None:
        try:
            style_elem = paragraph.style.element
            ancestor_numPr = style_elem.find(f".//{qn('w:numPr')}")
            if ancestor_numPr is not None:
                # Treat any styled list as unordered with ilvl=0 by default;
                # the renderer's smart merging handles consecutive items.
                return 2000, 0
        except Exception:  # pragma: no cover - defensive
            pass

    return None


def _is_ordered_list(doc: DocxDocument, num_id: int) -> bool:
    """Inspect numbering.xml to determine whether the list is ordered."""
    # Synthetic IDs from style-name fallback.
    if num_id == 1000:
        return True
    if num_id == 2000:
        return False
    numbering = getattr(doc.part, "numbering_part", None)
    if numbering is None:
        return False
    root = numbering.element
    num = root.find(f"{qn('w:num')}[@{qn('w:numId')}='{num_id}']")
    if num is None:
        return False
    abs_ref = num.find(qn("w:abstractNumId"))
    if abs_ref is None:
        return False
    abs_id = abs_ref.get(qn("w:val"))
    abs_def = root.find(f"{qn('w:abstractNum')}[@{qn('w:abstractNumId')}='{abs_id}']")
    if abs_def is None:
        return False
    lvl = abs_def.find(f"{qn('w:lvl')}[@{qn('w:ilvl')}='0']")
    if lvl is None:
        return False
    numFmt = lvl.find(qn("w:numFmt"))
    if numFmt is None:
        return False
    fmt = (numFmt.get(qn("w:val")) or "").lower()
    return fmt not in ("bullet", "none")


def _flush_list(buffer: list[tuple[int, ListItemNode]], doc, num_id: int) -> ListNode:
    """Turn a flat (ilvl, item) buffer into a nested ListNode tree."""
    # Build nested tree using a stack.
    roots: List[ListItemNode] = []
    stack: List[tuple[int, ListItemNode]] = []
    for ilvl, item in buffer:
        item.level = ilvl
        while stack and stack[-1][0] >= ilvl:
            stack.pop()
        if not stack:
            roots.append(item)
        else:
            stack[-1][1].sub_items.append(item)
        stack.append((ilvl, item))
    return ListNode(ordered=_is_ordered_list(doc, num_id), items=roots)


# ---------------------------------------------------------------------------
# Hyperlink, equation, citation detection
# ---------------------------------------------------------------------------

def _extract_hyperlinks(paragraph: Paragraph, doc: DocxDocument) -> List[HyperlinkNode]:
    out: List[HyperlinkNode] = []
    for hlink in paragraph._p.findall(qn("w:hyperlink")):
        rid = hlink.get(qn("r:id"))
        url = ""
        if rid is not None:
            try:
                url = paragraph.part.rels[rid].target_ref or ""
            except KeyError:
                url = ""
        runs: List[RunNode] = []
        for r_elem in hlink.findall(qn("w:r")):
            text_parts: list[str] = []
            for t in r_elem.findall(qn("w:t")):
                text_parts.append(t.text or "")
            text = "".join(text_parts)
            if text:
                runs.append(RunNode(text=text))
        if url and runs:
            out.append(HyperlinkNode(url=url, runs=runs))
    return out


def _extract_equations(paragraph: Paragraph) -> List[EquationNode]:
    """Find OMML oMath / oMathPara elements inside a paragraph."""
    out: List[EquationNode] = []
    # Display equations (oMathPara) — block-level
    for math_para in paragraph._p.findall(qn("m:oMathPara")):
        xml = etree.tostring(math_para, encoding="unicode")
        out.append(EquationNode(omml_xml=xml, display=True))
    # Inline equations — exclude those nested inside oMathPara already captured.
    for math in paragraph._p.findall(qn("m:oMath")):
        if math.getparent().tag == qn("m:oMathPara"):
            continue
        xml = etree.tostring(math, encoding="unicode")
        out.append(EquationNode(omml_xml=xml, display=False))
    return out


def _citation_tag(instr_text: str) -> Optional[str]:
    """Pull the source tag out of a ``CITATION <Tag> \\l 1033`` field code."""
    tokens = instr_text.split()
    # tokens[0] == 'CITATION'; the tag is the next non-switch token.
    for tok in tokens[1:]:
        if tok.startswith("\\"):
            break
        return tok
    return None


def _extract_citations(paragraph: Paragraph, bib=None) -> List[CitationNode]:
    """Pull Word citation fields out of a paragraph.

    Word emits citations as a sequence of w:r elements at the paragraph level:

        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText> CITATION ... </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:t>[1]</w:t></w:r>           ← display text lives here
        <w:r><w:fldChar w:fldCharType="end"/></w:r>

    We find each w:instrText containing CITATION, walk *up to its parent
    w:r*, then iterate forward over paragraph-level siblings (other w:r
    elements) until we hit the matching fldChar end, collecting w:t text
    after a fldChar separate marker.
    """
    out: List[CitationNode] = []
    for instr in paragraph._p.findall(f".//{qn('w:instrText')}"):
        if not instr.text or "CITATION" not in instr.text.upper():
            continue

        # The w:instrText is inside a w:r; walk forward across paragraph siblings.
        run_elem = instr.getparent()
        if run_elem is None:
            continue
        sibling = run_elem.getnext()
        display = ""
        seen_separator = False
        while sibling is not None:
            tag = etree.QName(sibling).localname
            if tag == "r":
                fld = sibling.find(qn("w:fldChar"))
                if fld is not None:
                    fc_type = fld.get(qn("w:fldCharType"))
                    if fc_type == "separate":
                        seen_separator = True
                    elif fc_type == "end":
                        break
                elif seen_separator:
                    for t in sibling.findall(qn("w:t")):
                        if t.text:
                            display += t.text
            sibling = sibling.getnext()

        raw = etree.tostring(instr, encoding="unicode")
        # Resolve the Word source tag to a BibTeX key when a bibliography was
        # extracted; source_id carries the key so the renderer emits \cite{}.
        tag = _citation_tag(instr.text or "")
        bib_key = bib.key_for_tag(tag) if bib is not None else None
        out.append(
            CitationNode(display_text=display or "?", raw_xml=raw, source_id=bib_key)
        )
    return out


# ---------------------------------------------------------------------------
# Footnotes
# ---------------------------------------------------------------------------

def _load_footnotes(doc: DocxDocument) -> dict[str, List[RunNode]]:
    """Parse word/footnotes.xml and return {id: runs}."""
    out: dict[str, List[RunNode]] = {}
    part = getattr(doc.part, "footnotes_part", None) or getattr(
        doc.part.package, "footnotes_part", None
    )
    if part is None:
        # python-docx doesn't expose footnotes natively; reach through package.
        for rel in doc.part.rels.values():
            if rel.reltype.endswith("/footnotes"):
                part = rel.target_part
                break
    if part is None:
        return out
    root = etree.fromstring(part.blob) if hasattr(part, "blob") else part.element
    for fn in root.findall(qn("w:footnote")):
        fn_id = fn.get(qn("w:id"))
        if fn_id in (None, "-1", "0"):
            continue
        runs: List[RunNode] = []
        for t in fn.findall(f".//{qn('w:t')}"):
            if t.text:
                runs.append(RunNode(text=t.text))
        out[fn_id] = merge_runs(runs)
    return out


def _extract_footnote_refs(paragraph: Paragraph, footnotes: dict) -> List[FootnoteNode]:
    out: List[FootnoteNode] = []
    for ref in paragraph._p.findall(f".//{qn('w:footnoteReference')}"):
        fn_id = ref.get(qn("w:id"))
        runs = footnotes.get(fn_id, [])
        if runs:
            out.append(FootnoteNode(runs=runs))
    return out


# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------

def _extract_images(
    paragraph: Paragraph,
    docx_zip: zipfile.ZipFile,
    counter: list[int],
    warnings: list[str],
) -> List[ImageNode]:
    out: List[ImageNode] = []
    blips = paragraph._p.findall(f".//{qn('a:blip')}")
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        if not rid:
            # r:link → externally linked image (URL on the author's machine).
            # We can't fetch it from the .docx zip; warn so the user knows the
            # figure went missing instead of silently dropping it.
            link_rid = blip.get(qn("r:link"))
            if link_rid:
                try:
                    target = paragraph.part.rels[link_rid].target_ref
                except KeyError:
                    target = "(unknown URL)"
                warnings.append(
                    f"External image link skipped (Word's r:link is not embedded "
                    f"in the .docx): {target}. Re-insert the image with "
                    f"'Insert as Embedded' in Word, then re-convert."
                )
            continue
        try:
            rel = paragraph.part.rels[rid]
        except KeyError:
            continue
        target = rel.target_ref
        # Compose the zip path: rel target is relative to word/
        zip_path = target if target.startswith("word/") else f"word/{target}"
        try:
            data = docx_zip.read(zip_path)
        except KeyError:
            continue
        counter[0] += 1
        ext = zip_path.rsplit(".", 1)[-1].lower()
        filename = f"figure_{counter[0]:03d}.{ext}"
        out.append(ImageNode(filename=filename, image_data=data))
    return out


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def _parse_cell(cell: _Cell, row_idx: int, col_idx: int, all_rows: list) -> TableCellNode:
    runs: List[RunNode] = []
    align = TextAlign.LEFT
    for p in cell.paragraphs:
        if align == TextAlign.LEFT:
            align = _paragraph_alignment(p)
        runs.extend(_build_runs(p))
    runs = merge_runs(runs)

    tcPr = cell._tc.find(qn("w:tcPr"))
    col_span = 1
    row_span = 1
    if tcPr is not None:
        gridSpan = tcPr.find(qn("w:gridSpan"))
        if gridSpan is not None:
            try:
                col_span = int(gridSpan.get(qn("w:val"), "1"))
            except ValueError:
                col_span = 1
        vMerge = tcPr.find(qn("w:vMerge"))
        if vMerge is not None and vMerge.get(qn("w:val")) == "restart":
            # Count how many subsequent rows have <w:vMerge/> (no val attr =
            # "continue") in the same column position. OOXML §17.4.84.
            row_span = 1
            for next_row in all_rows[row_idx + 1:]:
                if col_idx >= len(next_row.cells):
                    break
                next_tc = next_row.cells[col_idx]._tc
                next_tcPr = next_tc.find(qn("w:tcPr"))
                if next_tcPr is None:
                    break
                next_vMerge = next_tcPr.find(qn("w:vMerge"))
                if next_vMerge is None:
                    break
                # `continue` is encoded as <w:vMerge/> with no val OR val="continue".
                val = next_vMerge.get(qn("w:val"))
                if val and val != "continue":
                    break
                row_span += 1
    return TableCellNode(runs=runs, col_span=col_span, row_span=row_span, alignment=align)


def _parse_table(table: Table) -> TableNode:
    rows: List[TableRowNode] = []
    col_count = 0
    all_rows = list(table.rows)
    for idx, row in enumerate(all_rows):
        cells = [_parse_cell(c, idx, c_idx, all_rows) for c_idx, c in enumerate(row.cells)]
        col_count = max(col_count, sum(c.col_span for c in cells))
        # First row is treated as header unless the table only has one row.
        is_header = idx == 0 and len(table.rows) > 1
        rows.append(TableRowNode(cells=cells, is_header=is_header))
    # Aggregate column alignment from majority of cells in column 0..n
    col_alignments: List[TextAlign] = []
    for c_idx in range(col_count):
        votes: dict[TextAlign, int] = {}
        for row in rows:
            if c_idx < len(row.cells):
                a = row.cells[c_idx].alignment
                votes[a] = votes.get(a, 0) + 1
        if votes:
            col_alignments.append(max(votes.items(), key=lambda kv: kv[1])[0])
        else:
            col_alignments.append(TextAlign.LEFT)
    return TableNode(rows=rows, col_alignments=col_alignments)


# ---------------------------------------------------------------------------
# Main entrypoint
# ---------------------------------------------------------------------------

_MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024  # 200 MB hard ceiling
_MAX_UNCOMPRESSED_PER_FILE = 100 * 1024 * 1024  # 100 MB per entry
_MAX_ENTRIES = 5_000


def _assert_zip_safe(docx_bytes: bytes) -> None:
    """Reject zip bombs before any XML hits the parser.

    A .docx is a zip archive. A maliciously crafted 20 MB upload can carry a
    document.xml that decompresses to multiple gigabytes — enough to OOM-kill
    the worker the instant lxml tries to load it. We inspect the central
    directory (cheap; no decompression) and refuse files whose declared
    uncompressed size, per-entry or aggregate, exceeds sane ceilings.
    """
    with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
        infos = zf.infolist()
        if len(infos) > _MAX_ENTRIES:
            raise ValueError(
                f"Refusing zip with {len(infos)} entries (max {_MAX_ENTRIES})"
            )
        total = 0
        for info in infos:
            if info.file_size > _MAX_UNCOMPRESSED_PER_FILE:
                raise ValueError(
                    f"Refusing zip entry '{info.filename}' "
                    f"({info.file_size:,} B uncompressed; cap "
                    f"{_MAX_UNCOMPRESSED_PER_FILE:,})"
                )
            total += info.file_size
            if total > _MAX_UNCOMPRESSED_BYTES:
                raise ValueError(
                    f"Refusing zip: cumulative uncompressed payload exceeds "
                    f"{_MAX_UNCOMPRESSED_BYTES:,} B (possible zip bomb)"
                )


def parse_docx(docx_bytes: bytes) -> IRDocument:
    """Parse a .docx file's bytes into an IRDocument.

    Raises:
        ValueError: when the bytes are not a valid OOXML zip, or when the
            zip would decompress beyond the safety ceiling.
    """
    if not docx_bytes.startswith(b"PK\x03\x04"):
        raise ValueError("Not a valid .docx (missing PK zip magic bytes)")

    _assert_zip_safe(docx_bytes)

    doc: DocxDocument = Document(BytesIO(docx_bytes))
    docx_zip = zipfile.ZipFile(BytesIO(docx_bytes))

    warnings: List[str] = []
    nodes: List[IRNode] = []
    title: Optional[str] = None
    footnotes = _load_footnotes(doc)
    img_counter = [0]

    # v2 BibTeX: extract Word's managed Sources part once up front so each
    # in-text CITATION field can resolve to a \cite key.
    bib = extract_bibliography(docx_bytes)

    # Buffer for collecting consecutive list-item paragraphs.
    list_buffer: list[tuple[int, ListItemNode]] = []
    list_num_id: Optional[int] = None

    def flush_list():
        nonlocal list_buffer, list_num_id
        if list_buffer and list_num_id is not None:
            nodes.append(_flush_list(list_buffer, doc, list_num_id))
        list_buffer = []
        list_num_id = None

    # Walk the body in document order so paragraphs and tables remain interleaved.
    body = doc.element.body
    for child in body.iterchildren():
        tag = etree.QName(child).localname

        if tag == "p":
            p = Paragraph(child, doc)

            # w:pageBreakBefore in paragraph properties → \newpage BEFORE this
            # paragraph. Insert the page break first, then continue processing
            # the paragraph's content normally.
            if _paragraph_breaks_before(p):
                flush_list()
                nodes.append(PageBreakNode())

            info = _list_info(p)
            if info is not None:
                num_id, ilvl = info
                if list_num_id is None:
                    list_num_id = num_id
                elif num_id != list_num_id:
                    flush_list()
                    list_num_id = num_id
                runs = _build_runs(p)
                list_buffer.append((ilvl, ListItemNode(runs=runs, level=ilvl)))
                continue

            flush_list()

            # Page break-only paragraph (inline <w:br w:type="page"/>)
            if _has_page_break(p) and not p.text.strip():
                nodes.append(PageBreakNode())
                continue

            # Heading
            level = _heading_level(p)
            runs = _build_runs(p)
            if level is not None and runs:
                hn = HeadingNode(level=level, runs=runs)
                nodes.append(hn)
                if title is None and level == 1:
                    title = "".join(r.text for r in runs).strip()
                continue

            # Block-level extractions that REPLACE a paragraph
            equations = _extract_equations(p)
            images = _extract_images(p, docx_zip, img_counter, warnings)
            hyperlinks = _extract_hyperlinks(p, doc)
            citations = _extract_citations(p, bib)
            footnote_refs = _extract_footnote_refs(p, footnotes)

            # If the paragraph contains a display equation by itself, prefer the
            # equation node.
            if equations and not runs and not images:
                nodes.extend(equations)
                continue

            # Image-only paragraph
            if images and not runs:
                nodes.extend(images)
                continue

            # Otherwise emit a paragraph; attach inline hyperlinks/citations as
            # follow-on nodes (renderer respects document order).
            if runs:
                nodes.append(ParagraphNode(runs=runs, alignment=_paragraph_alignment(p)))
            nodes.extend(hyperlinks)
            nodes.extend(citations)
            nodes.extend(equations)
            nodes.extend(footnote_refs)
            nodes.extend(images)

            if _has_page_break(p):
                nodes.append(PageBreakNode())

        elif tag == "tbl":
            flush_list()
            table = Table(child, doc)
            try:
                nodes.append(_parse_table(table))
            except Exception as e:  # pragma: no cover - defensive
                warnings.append(f"Failed to parse table: {e}")

        elif tag == "sectPr":
            # A body-level sectPr defines the FINAL section break of the
            # document. Mid-document section breaks are encoded inside the
            # last paragraph of the previous section via pPr/sectPr — we
            # handle those by treating nextPage / oddPage / evenPage as a
            # page break before the next paragraph.
            if _section_break_after(child):
                # Only emit if the previous node isn't already a PageBreakNode
                # (avoids duplicate \newpage at end of document).
                if not nodes or not isinstance(nodes[-1], PageBreakNode):
                    nodes.append(PageBreakNode())
            continue

    flush_list()

    bibtex = None if bib.is_empty else bib.to_bibtex()
    return IRDocument(title=title, nodes=nodes, warnings=warnings, bibtex=bibtex)
